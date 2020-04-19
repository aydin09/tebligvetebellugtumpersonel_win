from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.shared import Mm, Cm, Inches
from docx.shared import Length
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk
import locale
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

locale.setlocale(locale.LC_ALL, "")

def teblig_konu(event):
    liste1=liste.get(ACTIVE)

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
    
    mudur_yard_adi_soyadi.delete(0,END)
    
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)
                       
    vt1 = sqlite3.connect(liste1+'.sql')
    im1= vt1.cursor()
    im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligeden TEXT, tebligedengorevi TEXT,tebligyazitarih TEXT,
                                                     tebligyazisayisi TEXT, tebligyazikonu TEXT, tebligyeri TEXT, tebligtarih TEXT, tebligsaat TEXT)""")
    
    im1.execute("""SELECT * FROM  teblig""")
    rows = im1.fetchall()
    data_str = ""
    sf = "{}{}{}{}{}{}{}{}{}{}"
    for row in rows:
        data_str += sf.format(row[0],row[1],row[2],row[3],row[4],row[5],row[6],row[7],row[8],row[9])

        kaymakamlik.insert(END,row[0])
        okul_adi.insert(END,row[1])
        teblig_eden.insert(END,row[2])
        teblig_eden_gorev.insert(END,row[3])
        teblig_edilen_yazi_tarih.insert(END,row[4])
        teblig_edilen_yazi_sayisi.insert(END,row[5])
        teblig_edilen_yazi_konusu.insert(END,row[6])
        teblig_yeri.insert(END,row[7])
        teblig_tarih.insert(END,row[8])
        teblig_saat.insert(END,row[9])
        
def mudur_yard(event):
    liste2=liste1.get(ACTIVE)

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
    
    mudur_yard_adi_soyadi.delete(0,END)
    
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)

    if liste2=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Müdür Yardımcısı Listesi Boş!').pack()
    else:
        vt1 = sqlite3.connect(liste2+'.sql3')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS teblig(muduryardadisoyadi TEXT)""")
        im1.execute("""SELECT * FROM  teblig""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}"
        for row in rows:
            data_str += sf.format(row[0])

            mudur_yard_adi_soyadi.insert(END,row[0])
        
def personel(event):
    liste3=liste2.get(ACTIVE)

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
    
    mudur_yard_adi_soyadi.delete(0,END)
    
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)

    if liste3=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Personel Listesi Boş!').pack()

    else:                 
        vt1 = sqlite3.connect(liste3+'.sq3')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS teblig(personeladisoyadi TEXT, personelgorev TEXT)""")
        im1.execute("""SELECT * FROM  teblig""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}"
        for row in rows:
            data_str += sf.format(row[0],row[1])

            personel_adi_soyadi.insert(END,row[0])
            personel_gorev.insert(END,row[1])
                        
def kaydet_teblig_edilen_yazi_konu():
    kaymakamlik1=kaymakamlik.get()
    okul_adi1=okul_adi.get()
    teblig_eden1=teblig_eden.get()
    teblig_eden_gorev1=teblig_eden_gorev.get()
    teblig_edilen_yazi_tarih1=teblig_edilen_yazi_tarih.get()
    teblig_edilen_yazi_sayisi1=teblig_edilen_yazi_sayisi.get()
    teblig_edilen_yazi_konusu1=teblig_edilen_yazi_konusu.get()
    teblig_yeri1=teblig_yeri.get()
    teblig_tarih1=teblig_tarih.get()
    teblig_saat1=teblig_saat.get()
    
    if kaymakamlik1=="" or okul_adi1=="" or teblig_eden1=="" or teblig_eden_gorev1=="" or teblig_edilen_yazi_tarih1=="" or teblig_edilen_yazi_sayisi1=="" or teblig_edilen_yazi_konusu1=="" or teblig_yeri1=="" or teblig_tarih1=="" or teblig_saat1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Bilgileri eksiksiz giriniz!').pack()

    else:    
        kaymakamlik.delete(0,END)
        okul_adi.delete(0,END)
        teblig_eden.delete(0,END)
        teblig_eden_gorev.delete(0,END)
        teblig_edilen_yazi_tarih.delete(0,END)
        teblig_edilen_yazi_sayisi.delete(0,END)
        teblig_edilen_yazi_konusu.delete(0,END)
        teblig_yeri.delete(0,END)
        teblig_tarih.delete(0,END)
        teblig_saat.delete(0,END)
        
        mudur_yard_adi_soyadi.delete(0,END)
        
        personel_adi_soyadi.delete(0,END)
        personel_gorev.delete(0,END)

        if os.path.exists(teblig_edilen_yazi_konusu1+'.sql')== False:
            vt1 = sqlite3.connect(teblig_edilen_yazi_konusu1+'.sql')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligeden TEXT, tebligedengorevi TEXT,tebligyazitarih TEXT,
                                                     tebligyazisayisi TEXT, tebligyazikonu TEXT, tebligyeri TEXT, tebligtarih TEXT, tebligsaat TEXT)""")
            
            im1.execute("""INSERT INTO teblig VALUES  (?,?,?,?,?,?,?,?,?,?)""",(kaymakamlik1, okul_adi1, teblig_eden1, teblig_eden_gorev1,
                                                                                teblig_edilen_yazi_tarih1, teblig_edilen_yazi_sayisi1, teblig_edilen_yazi_konusu1,
                                                                                teblig_yeri1, teblig_tarih1, teblig_saat1,))
            vt1.commit()

            liste.delete(0,END)

            for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
                if i.endswith('.sql'):
                    liste.insert(END,i[0:-4])

        else:
            vt2 = sqlite3.connect(teblig_edilen_yazi_konusu1+'.sql')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligeden TEXT, tebligedengorevi TEXT,tebligyazitarih TEXT,
                                                     tebligyazisayisi TEXT, tebligyazikonu TEXT, tebligyeri TEXT, tebligtarih TEXT, tebligsaat TEXT)""")
            
            im2.execute("""UPDATE teblig SET kaymakamlık=?, okul=?, tebligeden=?, tebligedengorevi=?,tebligyazitarih=?, tebligyazisayisi=?,
                        tebligyazikonu=?, tebligyeri=?, tebligtarih=?, tebligsaat=?""", (kaymakamlik1, okul_adi1, teblig_eden1, teblig_eden_gorev1,
                                                                                teblig_edilen_yazi_tarih1, teblig_edilen_yazi_sayisi1, teblig_edilen_yazi_konusu1,
                                                                                teblig_yeri1, teblig_tarih1, teblig_saat1,))
            
            vt2.commit()

def kaydet_mudur_yard():
    mudur_yard_adi_soyadi1=mudur_yard_adi_soyadi.get()

    if mudur_yard_adi_soyadi1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Müdür Yardımcısının Adını ve Soyadını Giriniz!').pack()
        mudur_yard_adi_soyadi.delete(0,END)

    else:    
        kaymakamlik.delete(0,END)
        okul_adi.delete(0,END)
        teblig_eden.delete(0,END)
        teblig_eden_gorev.delete(0,END)
        teblig_edilen_yazi_tarih.delete(0,END)
        teblig_edilen_yazi_sayisi.delete(0,END)
        teblig_edilen_yazi_konusu.delete(0,END)
        teblig_yeri.delete(0,END)
        teblig_tarih.delete(0,END)
        teblig_saat.delete(0,END)
        
        mudur_yard_adi_soyadi.delete(0,END)
        
        personel_adi_soyadi.delete(0,END)
        personel_gorev.delete(0,END)

        if os.path.exists(mudur_yard_adi_soyadi1+'.sql3')== False:
            vt1 = sqlite3.connect(mudur_yard_adi_soyadi1+'.sql3')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS teblig(muduryardadisoyadi TEXT)""")
            
            im1.execute("""INSERT INTO teblig VALUES  (?)""",(mudur_yard_adi_soyadi1,))
            vt1.commit()

            liste1.delete(0,END)

            for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
                if i.endswith('.sql3'):
                    liste1.insert(END,i[0:-5])

        else:
            vt2 = sqlite3.connect(mudur_yard_adi_soyadi1+'.sql3')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS teblig(muduryardadisoyadi TEXT)""")
            
            im2.execute("""UPDATE teblig SET muduryardadisoyadi=?""", (mudur_yard_adi_soyadi1,))
            
            vt2.commit()

def kaydet_personel():
    personel_adi_soyadi1=personel_adi_soyadi.get()
    personel_gorev1=personel_gorev.get()

    if personel_adi_soyadi1=="" or personel_gorev1=="":
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Personel Bilgilerini Tam Olarak Giriniz!').pack()

    else:    
        kaymakamlik.delete(0,END)
        okul_adi.delete(0,END)
        teblig_eden.delete(0,END)
        teblig_eden_gorev.delete(0,END)
        teblig_edilen_yazi_tarih.delete(0,END)
        teblig_edilen_yazi_sayisi.delete(0,END)
        teblig_edilen_yazi_konusu.delete(0,END)
        teblig_yeri.delete(0,END)
        teblig_tarih.delete(0,END)
        teblig_saat.delete(0,END)
        
        mudur_yard_adi_soyadi.delete(0,END)
        
        personel_adi_soyadi.delete(0,END)
        personel_gorev.delete(0,END)

        if os.path.exists(personel_adi_soyadi1+'.sq3')== False:
            vt1 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS teblig(personeladisoyadi TEXT, personelgorev TEXT)""")
            
            im1.execute("""INSERT INTO teblig VALUES  (?,?)""",(personel_adi_soyadi1, personel_gorev1,))
            vt1.commit()

            liste2.delete(0,END)

            for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
                if i.endswith('.sq3'):
                    liste2.insert(END,i[0:-4])

        else:
            vt2 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS teblig(personeladisoyadi TEXT, personelgorev TEXT)""")
            
            im2.execute("""UPDATE teblig SET personeladisoyadi=?, personelgorev=?""", (personel_adi_soyadi1, personel_gorev1,))
            
            vt2.commit()
       
def cikti():
    mud_yardimcisi_olan=[]

    personel_olan=[]
    
    kaymakamlik1=kaymakamlik.get()
    okul_adi1=okul_adi.get()
    teblig_eden1=teblig_eden.get()
    teblig_eden_gorev1=teblig_eden_gorev.get()
    teblig_edilen_yazi_tarih1=teblig_edilen_yazi_tarih.get()
    teblig_edilen_yazi_sayisi1=teblig_edilen_yazi_sayisi.get()
    teblig_edilen_yazi_konusu1=teblig_edilen_yazi_konusu.get()
    teblig_yeri1=teblig_yeri.get()
    teblig_tarih1=teblig_tarih.get()
    teblig_saat1=teblig_saat.get()

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
        if i.endswith('.sq3'):
            personel_olan.append(i[0:-4])

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):    
        if i.endswith('.sql3'):
            mud_yardimcisi_olan.append(i[0:-5])

    data4 = []
 
    while True:
        for s in range(0,len(mud_yardimcisi_olan)):
            
            vt4= sqlite3.connect(mud_yardimcisi_olan[s]+'.sql3')
            im4= vt4.cursor()
            im4.execute("""SELECT * FROM teblig""")
            rows4 = im4.fetchall()
            
            for row4 in rows4:
                data4.append(row4)

        break

    data3 = []
 
    while True:
        for s in range(0,len(personel_olan)):
            
            vt3= sqlite3.connect(personel_olan[s]+'.sq3')
            im3= vt3.cursor()
            im3.execute(""" SELECT * FROM teblig""")
            rows3 = im3.fetchall()
            
            for row3 in rows3:
                data3.append(row3)

        break
            
    if mud_yardimcisi_olan==[] or personel_olan==[] or kaymakamlik1=="" or okul_adi1=="" or teblig_eden1=="" or teblig_eden_gorev1=="" or teblig_edilen_yazi_tarih1=="" or teblig_edilen_yazi_sayisi1=="" or teblig_edilen_yazi_konusu1=="" or teblig_yeri1=="" or teblig_tarih1=="" or teblig_saat1=="" :
        uyari=Toplevel()
        uyari.resizable(width=FALSE ,height=FALSE)
        img=PhotoImage(file='teblig.png')
        uyari.tk.call('wm','iconphoto',uyari._w,img)
        Label(uyari, text ='Tebliğ Edilen Yazı Bilgilerini Eksiksiz Giriniz! veya\nMüdür Yardımcısının Adını ve Soyadını Girip Kaydediniz! veya\n Personel Bilgilerini Girip Kaydediniz!').pack()

    else:    
        vt1 = sqlite3.connect(teblig_edilen_yazi_konusu1+'.sql')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS teblig(kaymakamlık TEXT, okul TEXT, tebligeden TEXT, tebligedengorevi TEXT,tebligyazitarih TEXT,
                                                     tebligyazisayisi TEXT, tebligyazikonu TEXT, tebligyeri TEXT, tebligtarih TEXT, tebligsaat TEXT)""")
        im1.execute("""SELECT * FROM  teblig""")
        rows = im1.fetchall()
        data_str = ""
        sf = "{}{}{}{}{}{}{}{}{}{}"
        for row in rows:
            data_str += sf.format(row[0],row[1],row[2],row[3],row[4],row[5],row[6],row[7],row[8],row[9])
        vt1.commit()        
            
        document = Document()
        section = document.sections[0]
        
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
        
        for section in document.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(297) 
            section.page_height = Mm(210)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        table = document.add_table(rows=1, cols=1,style = 'Table Grid')
        table.columns[0].width = Cm(12.90)
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("T.C.\n"+row[0]+"\n"+row[1]+"\n"+"TEBLİĞ VE TEBELLÜĞ BELGESİ").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(14)

        paragraph = document.add_paragraph()

        table = document.add_table(rows=1, cols=1,style = 'Table Grid')
        table.columns[0].width = Cm(12.90)
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("TEBLİĞ EDİLEN YAZININ").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        table = document.add_table(rows=2, cols=4,style = 'Table Grid')
        table.columns[0].width = Cm(0.9)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        table.columns[3].width = Cm(4)
      
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("S.NO").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(6)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run("TARİHİ").bold = True
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run("SAYISI").bold = True
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,3)
        table.cell(0,3).paragraphs[0].add_run("KONUSU").bold = True
        table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,3).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("1").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(row[4])
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,2)
        table.cell(1,2).paragraphs[0].add_run(row[5])
        table.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,3)
        table.cell(1,3).paragraphs[0].add_run(row[6])
        table.cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,3).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        paragraph = document.add_paragraph()

        table = document.add_table(rows=2, cols=3,style = 'Table Grid')
        table.columns[0].width = Cm(4.9)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("TEBLİĞİN YERİ").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run("TEBLİĞİN TARİHİ").bold = True
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run("TEBLİĞ SAATİ").bold = True
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run(row[7])
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(row[8])
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,2)
        table.cell(1,2).paragraphs[0].add_run(row[9])
        table.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        paragraph = document.add_paragraph()

        table = document.add_table(rows=1, cols=1,style = 'Table Grid')
        table.columns[0].width = Cm(12.90)
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("TEBLİĞ EDEN").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        table = document.add_table(rows=2, cols=4,style = 'Table Grid')
        table.columns[0].width = Cm(0.9)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        table.columns[3].width = Cm(4)
      
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("S.NO").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(6)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run("ADI SOYADI").bold = True
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run("GÖREVİ").bold = True
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,2)
        table.cell(0,3).paragraphs[0].add_run("İMZA").bold = True
        table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,3).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,0)
        table.cell(1,0).paragraphs[0].add_run("1").bold = True
        table.cell(1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,1)
        table.cell(1,1).paragraphs[0].add_run(row[2])
        table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(1,2)
        table.cell(1,2).paragraphs[0].add_run(row[3])
        table.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(1,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        paragraph = document.add_paragraph()

        table = document.add_table(rows=2, cols=1,style = 'Table Grid')
        table.columns[0].width = Cm(12.90)
      
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("TEBELLÜĞ EDEN").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(12)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run("TARİH SAYI VE KONUSU BELİRTİLEN\n\nYAZIYI OKUDUM VE BİLGİ EDİNDİM.").bold = True
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        table = document.add_table(rows=1, cols=4,style = 'Table Grid')
        table.columns[0].width = Cm(0.9)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        table.columns[3].width = Cm(4)
      
        cell = table.cell(0,0)
        table.cell(0,0).paragraphs[0].add_run("S.NO").bold = True
        table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,0).paragraphs[0].runs[0].font
        cell_font.size = Pt(6)

        cell = table.cell(0,1)
        table.cell(0,1).paragraphs[0].add_run("ADI SOYADI").bold = True
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,1).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,2)
        table.cell(0,2).paragraphs[0].add_run("GÖREVİ").bold = True
        table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,2).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        cell = table.cell(0,3)
        table.cell(0,3).paragraphs[0].add_run("İMZA").bold = True
        table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_font = table.cell(0,3).paragraphs[0].runs[0].font
        cell_font.size = Pt(10)

        table = document.add_table(rows=len(mud_yardimcisi_olan)+len(personel_olan), cols=4,style = 'Table Grid')
        table.columns[0].width = Cm(0.9)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(4)
        table.columns[3].width = Cm(4)

        for s in range(0,len(mud_yardimcisi_olan)):
            cell = table.cell(s,0)
            cell.text =str(int(s)+1)
            table.cell(s,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_font = table.cell(s,0).paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            
            cell = table.cell(s,1)
            cell.text =data4[s][0]
            table.cell(s,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell_font = table.cell(s,1).paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            
            cell = table.cell(s,2)
            cell.text ="Müdür Yardımcısı"
            table.cell(s,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell_font = table.cell(s,2).paragraphs[0].runs[0].font
            cell_font.size = Pt(10)

        for s1 in range(0,len(personel_olan)):
            cell = table.cell(int(s1)+len(mud_yardimcisi_olan),0)
            cell.text =str(int(s1)+len(mud_yardimcisi_olan)+1)
            table.cell(s1+len(mud_yardimcisi_olan),0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_font = table.cell(int(s1)+1,0).paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            
            cell = table.cell(int(s1)+len(mud_yardimcisi_olan),1)
            cell.text =data3[s1][0]
            table.cell(len(mud_yardimcisi_olan),1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell_font = table.cell(int(s1)+1,1).paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            
            cell = table.cell(int(s1)+len(mud_yardimcisi_olan),2)
            cell.text =data3[s1][1]
            table.cell(len(mud_yardimcisi_olan),2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell_font = table.cell(int(s1)+len(mud_yardimcisi_olan),2).paragraphs[0].runs[0].font
            cell_font.size = Pt(10)
                
        document.save('tebligtebellug.docx')

        os.startfile("tebligtebellug.docx")

def sil_teblig_edilen():
    data_sil=liste.get(ACTIVE)

    os.remove(data_sil+".sql")

    liste.delete(0,END)

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
        if i.endswith('.sql'):
            liste.insert(END,i[0:-4])

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
        
    mudur_yard_adi_soyadi.delete(0,END)
        
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)

def sil_mudur_yard():
    data_sil=liste1.get(ACTIVE)

    os.remove(data_sil+".sql3")

    liste1.delete(0,END)

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
        if i.endswith('.sql3'):
            liste1.insert(END,i[0:-5])

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
        
    mudur_yard_adi_soyadi.delete(0,END)
        
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)

def sil_personel():
    data_sil=liste2.get(ACTIVE)

    os.remove(data_sil+".sq3")

    liste2.delete(0,END)

    for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
        if i.endswith('.sq3'):
            liste2.insert(END,i[0:-4])

    kaymakamlik.delete(0,END)
    okul_adi.delete(0,END)
    teblig_eden.delete(0,END)
    teblig_eden_gorev.delete(0,END)
    teblig_edilen_yazi_tarih.delete(0,END)
    teblig_edilen_yazi_sayisi.delete(0,END)
    teblig_edilen_yazi_konusu.delete(0,END)
    teblig_yeri.delete(0,END)
    teblig_tarih.delete(0,END)
    teblig_saat.delete(0,END)
        
    mudur_yard_adi_soyadi.delete(0,END)
        
    personel_adi_soyadi.delete(0,END)
    personel_gorev.delete(0,END)

root = Tk()
root.title("Tebliğ ve Tebellüğ Belgesi Programı")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='teblig.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

kaymakamlik = ttk.Entry(mainframe, width =40)
kaymakamlik.grid(column = 2, row = 0)

okul_adi = ttk.Entry(mainframe, width =40)
okul_adi.grid(column = 2, row = 1)

teblig_eden = ttk.Entry(mainframe, width =40)
teblig_eden.grid(column = 2, row = 2)

teblig_eden_gorev = ttk.Entry(mainframe, width =40)
teblig_eden_gorev.grid(column = 2, row = 3)

teblig_edilen_yazi_tarih = ttk.Entry(mainframe, width =40)
teblig_edilen_yazi_tarih.grid(column = 2, row = 4)

teblig_edilen_yazi_sayisi = ttk.Entry(mainframe, width =40)
teblig_edilen_yazi_sayisi.grid(column = 2, row = 5)

teblig_edilen_yazi_konusu = ttk.Entry(mainframe, width =40)
teblig_edilen_yazi_konusu.grid(column = 2, row = 6)

teblig_yeri = ttk.Entry(mainframe, width =40)
teblig_yeri.grid(column = 2, row = 7)

teblig_tarih = ttk.Entry(mainframe, width =40)
teblig_tarih.grid(column = 2, row = 8)

teblig_saat = ttk.Entry(mainframe, width =40)
teblig_saat.grid(column = 2, row = 9)

mudur_yard_adi_soyadi = ttk.Entry(mainframe, width =30)
mudur_yard_adi_soyadi.grid(column = 5, row = 33)

personel_adi_soyadi = ttk.Entry(mainframe, width =30)
personel_adi_soyadi.grid(column = 7, row = 33)

personel_gorev = ttk.Entry(mainframe, width =30)
personel_gorev.grid(column = 7, row = 37)

ttk.Label(mainframe, text ='KAYMAKAMLIK ADI').grid(column = 1, row = 0, sticky=W)
ttk.Label(mainframe, text ='OKULUN ADI').grid(column = 1, row = 1, sticky=W)
ttk.Label(mainframe, text ="TEBLİĞ EDENİN ADI SOYADI").grid(column = 1, row = 2, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDENİN GÖREVİ').grid(column = 1, row = 3, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDİLEN YAZININ TARİHİ').grid(column = 1, row = 4, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDİLEN YAZININ SAYISI').grid(column = 1, row = 5, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞ EDİLEN YAZININ KONUSU').grid(column = 1, row = 6, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN YERİ').grid(column = 1, row = 7, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN TARİHİ').grid(column = 1, row = 8, sticky=W)
ttk.Label(mainframe, text ='TEBLİĞİN SAATİ').grid(column = 1, row = 9, sticky=W)

ttk.Label(mainframe, text ='MÜDÜR YARDIMCISININ ADI SOYADI').grid(column = 5, row = 32)
ttk.Label(mainframe, text ='PERSONELİN ADI SOYADI').grid(column = 7, row = 32)
ttk.Label(mainframe, text ='PERSONELİN GÖREVİ').grid(column = 7, row = 35)

ttk.Label(mainframe, text ='TEBLİĞ EDİLEN YAZININ KONUSU').grid(column = 3, row=0)
ttk.Label(mainframe, text ='MÜDÜR YARDIMCISI LİSTESİ').grid(column = 5, row=0)
ttk.Label(mainframe, text ='PERSONEL LİSTESİ').grid(column = 7, row=0)

ttk.Label(mainframe, text ='').grid(column = 3, row=31)
ttk.Label(mainframe, text ='').grid(column = 7, row=38)
ttk.Label(mainframe, text ='').grid(column = 7, row=40)

liste = Listbox(mainframe,width=30)
liste.grid(column=3, row=1,rowspan=20,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",teblig_konu)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=4, row=1, rowspan=20,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

liste1 = Listbox(mainframe,width=20)
liste1.grid(column=5, row=1,rowspan=20,  sticky=(N,S,E,W))
liste1.bind("<Double-Button-1>",mudur_yard)

kaydirma1 = ttk.Scrollbar(mainframe, orient="vertical",command=liste1.yview)
kaydirma1.grid(column=6, row=1, rowspan=20,sticky='ns')

liste1.config(yscrollcommand=kaydirma1.set)
kaydirma1.config(command=liste1.yview)

liste2 = Listbox(mainframe,width=20)
liste2.grid(column=7, row=1,rowspan=20,  sticky=(N,S,E,W))
liste2.bind("<Double-Button-1>",personel)

kaydirma2 = ttk.Scrollbar(mainframe, orient="vertical",command=liste2.yview)
kaydirma2.grid(column=8, row=1, rowspan=20,sticky='ns')

liste2.config(yscrollcommand=kaydirma2.set)
kaydirma2.config(command=liste2.yview)

for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
    if i.endswith('.sql'):
        liste.insert(END,i[0:-4])

for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
    if i.endswith('.sql3'):
        liste1.insert(END,i[0:-5])

for i in sorted(os.listdir(os.getcwd()), key=locale.strxfrm):
    if i.endswith('.sq3'):
        liste2.insert(END,i[0:-4])

ttk.Button(mainframe, text='Tebliğ Edilen Yazının Konusunu\n              Kaydet/Güncelle',command= kaydet_teblig_edilen_yazi_konu).grid(column=3, row=32)
ttk.Button(mainframe, text='Sil', command= sil_teblig_edilen).grid(column=3, row=34)

ttk.Button(mainframe, text='Müdür Yardımcısını Kaydet/Güncelle',command= kaydet_mudur_yard).grid(column=5, row=35)
ttk.Button(mainframe, text='Sil', command= sil_mudur_yard).grid(column=5, row=38)

ttk.Button(mainframe, text='Personeli Kaydet/Güncelle',command= kaydet_personel).grid(column=7, row=39)
ttk.Button(mainframe, text='Sil', command= sil_personel).grid(column=7, row=41)

ttk.Button(mainframe, text='Tebliğ ve Tebellüğ Ön İzleme', command = cikti).grid(column=2, row=32)

kaymakamlik.focus()

root.mainloop()    
